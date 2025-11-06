"""
Professional Data Cleaning Platform
Enterprise-grade data quality management with comprehensive web reports
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import requests
import time
import io
from datetime import datetime
import json
import base64
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
import tempfile
import os

# Page Configuration
st.set_page_config(
    page_title="DataClean Pro - Professional Platform",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Professional CSS with animations and transitions
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Roboto', sans-serif;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        min-height: 100vh;
        font-size: 18px !important;
    }
    
    .main {
        padding: 1rem;
        animation: fadeIn 0.8s ease-in;
        font-size: 20px !important;
    }
    
    /* Streamlit elements font size overrides */
    .stMarkdown, .stText, p, div, span {
        font-size: 20px !important;
    }
    
    h1 {
        font-size: 48px !important;
        font-weight: 700 !important;
    }
    
    h2 {
        font-size: 36px !important;
        font-weight: 600 !important;
    }
    
    h3 {
        font-size: 28px !important;
        font-weight: 500 !important;
    }
    
    h4 {
        font-size: 24px !important;
        font-weight: 500 !important;
    }
    
    .stButton > button {
        font-size: 18px !important;
        font-weight: 600 !important;
    }
    
    .stSelectbox > div > div, .stFileUploader, .stAlert {
        font-size: 18px !important;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    @keyframes slideInLeft {
        from { opacity: 0; transform: translateX(-50px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    @keyframes slideInRight {
        from { opacity: 0; transform: translateX(50px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    @keyframes scaleIn {
        from { opacity: 0; transform: scale(0.8); }
        to { opacity: 1; transform: scale(1); }
    }
    
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.05); }
        100% { transform: scale(1); }
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {visibility: hidden;}
    
    .header-section {
        background: linear-gradient(135deg, #2c3e50 0%, #3498db 50%, #9b59b6 100%);
        padding: 3rem 2rem;
        margin: -1rem -1rem 2rem -1rem;
        color: white;
        text-align: center;
        border-radius: 0 0 25px 25px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.3);
        position: relative;
        overflow: hidden;
    }
    
    .header-section::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(45deg, transparent, rgba(255,255,255,0.1), transparent);
        transform: rotate(45deg);
        animation: shine 3s infinite;
    }
    
    @keyframes shine {
        0% { transform: translateX(-100%) translateY(-100%) rotate(45deg); }
        100% { transform: translateX(100%) translateY(100%) rotate(45deg); }
    }
    
    .header-title {
        font-size: 4.5rem !important;
        font-weight: 700;
        margin-bottom: 0.5rem;
        background: linear-gradient(45deg, #fff, #ecf0f1);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        position: relative;
        z-index: 1;
        line-height: 1.2;
    }
    
    .header-subtitle {
        font-size: 2rem !important;
        opacity: 0.95;
        font-weight: 400;
        position: relative;
        z-index: 1;
        line-height: 1.3;
    }
    
    .metric-container {
        background: linear-gradient(135deg, #fff 0%, #f8f9fa 100%);
        border-radius: 15px;
        padding: 2rem;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
        border: none;
        border-left: 5px solid #3498db;
        transition: all 0.3s ease;
        animation: slideInLeft 0.6s ease-out;
        position: relative;
        overflow: hidden;
    }
    
    .metric-container::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 2px;
        background: linear-gradient(90deg, #3498db, #9b59b6, #e74c3c);
        transform: translateX(-100%);
        transition: transform 0.6s ease;
    }
    
    .metric-container:hover::before {
        transform: translateX(0);
    }
    
    .metric-container:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 35px rgba(0,0,0,0.15);
    }
    
    .metric-title {
        font-size: 1.3rem !important;
        color: #666;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 0.5rem;
        font-weight: 500;
    }
    
    .metric-value {
        font-size: 3.5rem !important;
        font-weight: 700;
        color: #2c3e50;
        background: linear-gradient(45deg, #2c3e50, #3498db);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        line-height: 1.1;
    }
    
    .chart-container {
        background: linear-gradient(135deg, #fff 0%, #f8f9fa 100%);
        border-radius: 20px;
        padding: 2.5rem;
        margin: 1.5rem 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
        animation: scaleIn 0.7s ease-out;
        border: 1px solid rgba(255,255,255,0.8);
    }
    
    .chart-container:hover {
        transform: translateY(-3px);
        box-shadow: 0 15px 40px rgba(0,0,0,0.15);
    }
    
    .chart-title {
        font-size: 2.2rem !important;
        font-weight: 600;
        margin-bottom: 1.5rem;
        color: #2c3e50;
        text-align: center;
        position: relative;
    }
    
    .chart-title::after {
        content: '';
        position: absolute;
        bottom: -8px;
        left: 50%;
        transform: translateX(-50%);
        width: 60px;
        height: 3px;
        background: linear-gradient(90deg, #3498db, #9b59b6);
        border-radius: 2px;
    }
    
    .upload-container {
        background: linear-gradient(135deg, #fff 0%, #f8f9fa 100%);
        border: 3px dashed #bdc3c7;
        border-radius: 20px;
        padding: 3rem;
        text-align: center;
        margin: 2rem 0;
        transition: all 0.4s ease;
        position: relative;
        overflow: hidden;
        animation: slideInRight 0.8s ease-out;
    }
    
    .upload-container::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(52, 152, 219, 0.1), transparent);
        transition: left 0.6s ease;
    }
    
    .upload-container:hover {
        border-color: #3498db;
        background: linear-gradient(135deg, #f8f9fa 0%, #ecf0f1 100%);
        transform: scale(1.02);
        box-shadow: 0 10px 30px rgba(52, 152, 219, 0.2);
    }
    
    .upload-container:hover::before {
        left: 100%;
    }
    
    .status-success {
        background: linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(46, 204, 113, 0.3);
        animation: slideInLeft 0.5s ease-out;
        position: relative;
        overflow: hidden;
    }
    
    .status-success::before {
        content: 'OK';
        position: absolute;
        right: 20px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 2rem;
        opacity: 0.3;
    }
    
    .status-warning {
        background: linear-gradient(135deg, #f39c12 0%, #e67e22 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(243, 156, 18, 0.3);
        animation: pulse 2s infinite;
    }
    
    .status-error {
        background: linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
        box-shadow: 0 8px 25px rgba(231, 76, 60, 0.3);
        animation: slideInRight 0.5s ease-out;
    }
    
    .progress-container {
        background: linear-gradient(135deg, #fff 0%, #f8f9fa 100%);
        border-radius: 20px;
        padding: 3rem 2rem;
        margin: 2rem 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        animation: fadeIn 1s ease-out;
        border: 1px solid rgba(255,255,255,0.8);
    }
    
    .step-indicator {
        display: flex;
        justify-content: space-between;
        margin: 2rem 0;
        position: relative;
        padding: 0 25px;
    }
    
    .step-indicator::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 25px;
        right: 25px;
        height: 4px;
        background: linear-gradient(90deg, #ecf0f1, #bdc3c7);
        z-index: 1;
        border-radius: 2px;
    }
    
    .step {
        background: linear-gradient(135deg, #ecf0f1 0%, #bdc3c7 100%);
        color: #95a5a6;
        width: 70px;
        height: 70px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 600;
        position: relative;
        z-index: 2;
        transition: all 0.4s ease;
        cursor: pointer;
        border: 3px solid white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        font-size: 1.5rem !important;
    }
    
    .step:hover {
        transform: scale(1.1);
        box-shadow: 0 6px 20px rgba(0,0,0,0.2);
    }
    
    .step.active {
        background: linear-gradient(135deg, #3498db 0%, #2980b9 100%);
        color: white;
        transform: scale(1.15);
        box-shadow: 0 8px 25px rgba(52, 152, 219, 0.4);
        animation: pulse 2s infinite;
    }
    
    .step.completed {
        background: linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);
        color: white;
        transform: scale(1.05);
        box-shadow: 0 6px 20px rgba(46, 204, 113, 0.3);
    }
    
    .step.clickable {
        cursor: pointer;
        position: relative;
    }
    
    .step.clickable::after {
        content: '';
        position: absolute;
        top: -5px;
        left: -5px;
        right: -5px;
        bottom: -5px;
        border-radius: 50%;
        border: 2px solid transparent;
        background: linear-gradient(45deg, #3498db, #9b59b6) border-box;
        mask: linear-gradient(#fff 0 0) padding-box, linear-gradient(#fff 0 0);
        mask-composite: exclude;
        opacity: 0;
        transition: opacity 0.3s ease;
    }
    
    .step.clickable:hover::after {
        opacity: 1;
        animation: rotate 2s linear infinite;
    }
    
    @keyframes rotate {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }
    
    .report-section {
        background: linear-gradient(135deg, #fff 0%, #f8f9fa 100%);
        border-radius: 20px;
        padding: 2.5rem;
        margin: 1.5rem 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        animation: slideInLeft 0.6s ease-out;
        border: 1px solid rgba(255,255,255,0.8);
        position: relative;
        overflow: hidden;
    }
    
    .report-section::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: linear-gradient(90deg, #3498db, #9b59b6, #e74c3c, #f39c12);
        background-size: 300% 100%;
        animation: gradientShift 3s ease infinite;
    }
    
    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    
    .section-header {
        font-size: 2.4rem !important;
        font-weight: 600;
        margin-bottom: 1.5rem;
        color: #2c3e50;
        border-bottom: 3px solid #ecf0f1;
        padding-bottom: 0.8rem;
        position: relative;
        background: linear-gradient(45deg, #2c3e50, #3498db);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .data-preview {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 1px solid #e9ecef;
        transition: all 0.3s ease;
        animation: fadeIn 0.8s ease-out;
    }
    
    .data-preview:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
    }
    
    .cleaning-step {
        background: linear-gradient(135deg, #e8f6f3 0%, #d5f4e6 100%);
        border-left: 5px solid #27ae60;
        padding: 1.5rem;
        margin: 1rem 0;
        border-radius: 0 12px 12px 0;
        transition: all 0.3s ease;
        animation: slideInRight 0.6s ease-out;
        position: relative;
        overflow: hidden;
    }
    
    .cleaning-step::before {
        content: 'DONE';
        position: absolute;
        right: 15px;
        top: 15px;
        color: #27ae60;
        font-weight: bold;
        font-size: 1.2rem;
        opacity: 0.5;
    }
    
    .cleaning-step:hover {
        transform: translateX(5px);
        box-shadow: 0 8px 25px rgba(39, 174, 96, 0.2);
    }
    
    .issue-item {
        background: linear-gradient(135deg, #fdf2e9 0%, #fbeee6 100%);
        border-left: 5px solid #e67e22;
        padding: 1.5rem;
        margin: 1rem 0;
        border-radius: 0 12px 12px 0;
        transition: all 0.3s ease;
        animation: slideInLeft 0.6s ease-out;
        position: relative;
    }
    
    .issue-item:hover {
        transform: translateX(-5px);
        box-shadow: 0 8px 25px rgba(230, 126, 34, 0.2);
    }
    
    .issue-high {
        border-left-color: #e74c3c;
        background: linear-gradient(135deg, #fdedec 0%, #fadbd8 100%);
    }
    
    .issue-medium {
        border-left-color: #f39c12;
        background: linear-gradient(135deg, #fdf2e9 0%, #fbeee6 100%);
    }
    
    .issue-low {
        border-left-color: #f1c40f;
        background: linear-gradient(135deg, #fffbea 0%, #fef9e7 100%);
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #3498db 0%, #2980b9 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 1rem 2.5rem;
        font-weight: 600;
        font-size: 18px !important;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(52, 152, 219, 0.3);
        position: relative;
        overflow: hidden;
        min-height: 50px;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
        transition: left 0.6s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(52, 152, 219, 0.4);
    }
    
    .stButton > button:hover::before {
        left: 100%;
    }
    
    .stDownloadButton > button {
        background: linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 1rem 2.5rem;
        font-weight: 600;
        font-size: 18px !important;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(46, 204, 113, 0.3);
        min-height: 50px;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(46, 204, 113, 0.4);
    }
    
    /* Spinner Animation */
    .stSpinner > div {
        border-color: #3498db !important;
    }
    
    /* Loading Animation */
    @keyframes loading {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .loading-spinner {
        border: 4px solid #f3f3f3;
        border-top: 4px solid #3498db;
        border-radius: 50%;
        width: 40px;
        height: 40px;
        animation: loading 1s linear infinite;
        margin: 20px auto;
    }
    
    /* Fade in animation for content */
    .fade-in {
        animation: fadeIn 0.8s ease-in;
    }
    
    /* Slide animations */
    .slide-in-left {
        animation: slideInLeft 0.6s ease-out;
    }
    
    .slide-in-right {
        animation: slideInRight 0.6s ease-out;
    }
    
    /* Scale animation */
    .scale-in {
        animation: scaleIn 0.7s ease-out;
    }
</style>
""", unsafe_allow_html=True)

# API Configuration (auto-detect 8003 → 8000 fallback)
API_BASE_URL = "http://localhost:8003"

# Session state initialization
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'file_id' not in st.session_state:
    st.session_state.file_id = None
if 'analysis_data' not in st.session_state:
    st.session_state.analysis_data = None
if 'cleaning_report' not in st.session_state:
    st.session_state.cleaning_report = None
if 'original_data' not in st.session_state:
    st.session_state.original_data = None
if 'cleaned_data' not in st.session_state:
    st.session_state.cleaned_data = None
if 'show_complete_report' not in st.session_state:
    st.session_state.show_complete_report = False


def check_api_connection():
    """Check if backend API is available; fallback to port 8000 if 8003 is down."""
    global API_BASE_URL
    try:
        response = requests.get(f"{API_BASE_URL}/health", timeout=2)
        if response.status_code == 200:
            return True
    except Exception:
        pass
    # Try fallback port 8000 (main.py based server)
    try:
        fallback = "http://localhost:8000"
        resp = requests.get(f"{fallback}/health", timeout=2)
        if resp.status_code == 200:
            API_BASE_URL = fallback
            return True
    except Exception:
        pass
    return False


def create_header():
    """Create professional header"""
    st.markdown("""
    <div class="header-section">
        <h1 class="header-title">DataClean Pro</h1>
        <p class="header-subtitle">
            Data Quality Management System
        </p>
    </div>
    """, unsafe_allow_html=True)


def generate_pdf_report(analysis_data, cleaning_report=None):
    """Generate comprehensive PDF report"""
    if not REPORTLAB_AVAILABLE:
        st.error(
            "PDF generation requires reportlab. Please install it: pip install reportlab")
        return None

    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors

        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_file.close()  # Close the file handle before using it

        # Create PDF document with better margins
        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=A4,
            topMargin=0.8*inch,
            bottomMargin=0.8*inch,
            leftMargin=1*inch,
            rightMargin=1*inch
        )

        styles = getSampleStyleSheet()
        story = []

        # Custom styles for better formatting
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=22,
            fontName='Helvetica-Bold',
            spaceAfter=24,
            spaceBefore=0,
            alignment=1,  # Center
            textColor=colors.HexColor('#2c3e50')
        )

        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=styles['Heading2'],
            fontSize=16,
            fontName='Helvetica-Bold',
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.HexColor('#34495e')
        )

        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica',
            spaceAfter=8,
            leading=14
        )

        # Title
        story.append(Paragraph(
            "Comprehensive Data Quality Analysis & Cleaning Report", title_style))
        story.append(Spacer(1, 20))

        # Report Information Box with better formatting
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        info_data = [
            ['Report Generated:', timestamp],
            ['System:', 'Autonomous Data Cleaning Platform'],
            ['Status:', 'Complete Analysis & Cleaning']
        ]

        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
            ('BORDER', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 25))

        # Executive Summary
        story.append(Paragraph("Executive Summary", header_style))

        # Dataset Overview
        total_rows = analysis_data.get('total_rows', 0)
        total_columns = analysis_data.get('total_columns', 0)
        total_issues = len(analysis_data.get('issues', []))

        summary_data = [['Metric', 'Value']]
        summary_data.append(
            ['Original Size', f'{total_rows:,} rows × {total_columns} columns'])
        summary_data.append(['Total Issues Identified', str(total_issues)])

        if cleaning_report:
            before_quality = cleaning_report.get(
                'before_quality', {}).get('quality_score', 0)
            after_quality = cleaning_report.get(
                'after_quality', {}).get('quality_score', 0)
            improvement = after_quality - before_quality
            before_issues = len(cleaning_report.get(
                'before_quality', {}).get('issues', []))
            after_issues = len(cleaning_report.get(
                'after_quality', {}).get('issues', []))

            summary_data.extend([
                ['Quality Score Before Cleaning', f'{before_quality:.1f}%'],
                ['Quality Score After Cleaning', f'{after_quality:.1f}%'],
                ['Quality Improvement', f'{improvement:+.1f}%'],
                ['Issues Resolved',
                    f'{before_issues - after_issues} out of {before_issues}'],
                ['Cleaning Success Rate',
                    f'{((before_issues - after_issues) / before_issues * 100) if before_issues > 0 else 100:.1f}%']
            ])

        summary_table = Table(summary_data, colWidths=[3*inch, 3*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1),
             [colors.white, colors.HexColor('#f8f9fa')])
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 25))

        # Original Data Quality Issues
        if analysis_data.get('issues'):
            story.append(
                Paragraph("Original Data Quality Issues", header_style))

            # Group issues by severity for better organization
            high_issues = [issue for issue in analysis_data['issues']
                           if issue.get('severity', 'medium') == 'high']
            medium_issues = [issue for issue in analysis_data['issues'] if issue.get(
                'severity', 'medium') == 'medium']
            low_issues = [issue for issue in analysis_data['issues']
                          if issue.get('severity', 'medium') == 'low']

            issues_data = [['Column', 'Issue Type',
                            'Count', 'Percentage', 'Severity']]

            # Add issues in order of severity
            for issue in high_issues + medium_issues + low_issues:
                issues_data.append([
                    issue.get('column', 'N/A'),
                    issue.get('type', 'Unknown').replace('_', ' ').title(),
                    str(issue.get('count', 0)),
                    f"{issue.get('percentage', 0):.1f}%",
                    issue.get('severity', 'Medium').title()
                ])

            # Create table with appropriate column widths
            issues_table = Table(issues_data, colWidths=[
                                 1.5*inch, 1.8*inch, 0.8*inch, 1*inch, 1*inch])
            issues_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e74c3c')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#f8f9fa')])
            ]))
            story.append(issues_table)
            story.append(Spacer(1, 25))

        # Cleaning Operations Applied (if cleaning was performed)
        if cleaning_report and cleaning_report.get('steps_applied'):
            story.append(PageBreak())
            story.append(
                Paragraph("Applied Cleaning Operations", header_style))

            steps = cleaning_report['steps_applied']
            for i, step in enumerate(steps, 1):
                if isinstance(step, dict):
                    step_title = f"{i}. {step.get('step', 'Cleaning Operation')}"
                    story.append(Paragraph(step_title, ParagraphStyle(
                        'StepTitle', parent=normal_style, fontName='Helvetica-Bold', fontSize=12
                    )))

                    if step.get('description'):
                        story.append(
                            Paragraph(f"Description: {step['description']}", normal_style))

                    if step.get('details'):
                        details = step['details']
                        if isinstance(details, list):
                            for detail in details:
                                story.append(
                                    Paragraph(f"• {detail}", normal_style))
                        else:
                            story.append(
                                Paragraph(f"• {details}", normal_style))
                else:
                    story.append(Paragraph(f"{i}. {step}", ParagraphStyle(
                        'StepTitle', parent=normal_style, fontName='Helvetica-Bold', fontSize=12
                    )))

                story.append(Spacer(1, 10))

        # Final Quality Assessment
        if cleaning_report:
            story.append(
                Paragraph("Post-Cleaning Quality Assessment", header_style))

            assessment_text = f"""
            <b>Quality Improvement Summary:</b><br/>
            • Initial Quality Score: {before_quality:.1f}%<br/>
            • Final Quality Score: {after_quality:.1f}%<br/>
            • Overall Improvement: {improvement:+.1f}%<br/>
            • Issues Successfully Resolved: {before_issues - after_issues}<br/>
            • Remaining Issues: {after_issues}<br/><br/>
            """

            if after_quality >= 95:
                assessment_text += "<b>Assessment:</b> Excellent - Data quality is now at optimal level for analysis and modeling."
            elif after_quality >= 85:
                assessment_text += "<b>Assessment:</b> Good - Data quality is acceptable for most analytical purposes."
            elif after_quality >= 70:
                assessment_text += "<b>Assessment:</b> Fair - Data quality has improved but may benefit from additional domain-specific cleaning."
            else:
                assessment_text += "<b>Assessment:</b> Needs Attention - Consider additional data sources or manual review for critical applications."

            story.append(Paragraph(assessment_text, normal_style))
            story.append(Spacer(1, 20))

        # Recommendations
        story.append(Paragraph("Recommendations & Next Steps", header_style))
        recommendations_text = """
        <b>Data Quality Recommendations:</b><br/>
        • Implement regular data quality monitoring for ongoing data ingestion<br/>
        • Consider implementing data validation rules at the source to prevent future quality issues<br/>
        • Establish data governance policies for consistent data handling<br/>
        • Regular automated quality checks should be scheduled<br/>
        • Document data transformation rules for reproducibility<br/><br/>
        
        <b>Technical Recommendations:</b><br/>
        • Set up automated data validation pipelines<br/>
        • Implement exception handling for edge cases<br/>
        • Create data quality dashboards for monitoring<br/>
        • Establish backup and recovery procedures for cleaned data
        """
        story.append(Paragraph(recommendations_text, normal_style))
        story.append(Spacer(1, 20))

        # Footer
        footer_style = ParagraphStyle(
            'FooterStyle',
            parent=styles['Normal'],
            fontSize=9,
            alignment=1,
            textColor=colors.HexColor('#7f8c8d')
        )
        story.append(Paragraph(
            "Generated by Autonomous Data Cleaning Platform | Data Quality Management", footer_style))

        doc.build(story)

        # Read the file content
        with open(temp_file.name, 'rb') as f:
            pdf_content = f.read()

        # Clean up
        os.unlink(temp_file.name)

        return pdf_content

    except ImportError:
        st.error("PDF generation requires reportlab library. Please install it.")
        return None
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None


def generate_word_report(analysis_data, cleaning_report=None):
    """Generate comprehensive Word document report with detailed analysis and cleaning results"""
    if not DOCX_AVAILABLE:
        st.error(
            "Word document generation is not available. python-docx package is not installed.")
        return None

    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(
            'Comprehensive Data Quality Analysis & Cleaning Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report Information
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        info_para = doc.add_paragraph()
        info_para.add_run('Report Generated: ').bold = True
        info_para.add_run(f'{timestamp}\n')
        info_para.add_run('System: ').bold = True
        info_para.add_run('AI-Powered Data Cleaning Platform\n')
        info_para.add_run('Status: ').bold = True
        info_para.add_run('Complete Analysis & Cleaning')

        doc.add_paragraph()  # Empty line

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)

        # Dataset Overview
        total_rows = analysis_data.get('total_rows', 0)
        total_columns = analysis_data.get('total_columns', 0)
        total_issues = len(analysis_data.get('issues', []))

        summary = doc.add_paragraph()
        summary.add_run('Dataset Overview:\n').bold = True
        summary.add_run(
            f'• Original Size: {total_rows:,} rows × {total_columns} columns\n')
        summary.add_run(f'• Total Issues Identified: {total_issues}\n')

        if cleaning_report:
            before_quality = cleaning_report.get(
                'before_quality', {}).get('quality_score', 0)
            after_quality = cleaning_report.get(
                'after_quality', {}).get('quality_score', 0)
            improvement = after_quality - before_quality
            before_issues = len(cleaning_report.get(
                'before_quality', {}).get('issues', []))
            after_issues = len(cleaning_report.get(
                'after_quality', {}).get('issues', []))

            summary.add_run(
                f'• Quality Score Before Cleaning: {before_quality:.1f}%\n')
            summary.add_run(
                f'• Quality Score After Cleaning: {after_quality:.1f}%\n')
            summary.add_run(f'• Quality Improvement: {improvement:+.1f}%\n')
            summary.add_run(
                f'• Issues Resolved: {before_issues - after_issues} out of {before_issues}\n')
            summary.add_run(
                f'• Cleaning Success Rate: {((before_issues - after_issues) / before_issues * 100) if before_issues > 0 else 100:.1f}%\n')

        # Original Data Quality Issues
        if analysis_data.get('issues'):
            doc.add_heading('Original Data Quality Issues', level=1)

            # Group issues by severity
            high_issues = [issue for issue in analysis_data['issues']
                           if issue.get('severity', 'medium') == 'high']
            medium_issues = [issue for issue in analysis_data['issues'] if issue.get(
                'severity', 'medium') == 'medium']
            low_issues = [issue for issue in analysis_data['issues']
                          if issue.get('severity', 'medium') == 'low']

            # Create comprehensive table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Column'
            hdr_cells[1].text = 'Issue Type'
            hdr_cells[2].text = 'Count'
            hdr_cells[3].text = 'Percentage'
            hdr_cells[4].text = 'Severity'
            hdr_cells[5].text = 'Description'

            # Make headers bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Data rows - high priority first
            for issue in high_issues + medium_issues + low_issues:
                row_cells = table.add_row().cells
                row_cells[0].text = issue.get('column', 'N/A')
                row_cells[1].text = issue.get(
                    'type', 'Unknown').replace('_', ' ').title()
                row_cells[2].text = str(issue.get('count', 0))
                row_cells[3].text = f"{issue.get('percentage', 0):.1f}%"
                row_cells[4].text = issue.get('severity', 'Medium').title()
                description = issue.get(
                    'description', 'No description available')
                row_cells[5].text = description[:100] + \
                    ('...' if len(description) > 100 else '')

        # Applied Cleaning Operations (if cleaning was performed)
        if cleaning_report and cleaning_report.get('steps_applied'):
            doc.add_page_break()
            doc.add_heading('Applied Cleaning Operations', level=1)

            steps = cleaning_report['steps_applied']
            for i, step in enumerate(steps, 1):
                step_para = doc.add_paragraph()
                step_para.add_run(f'{i}. ').bold = True

                if isinstance(step, dict):
                    step_para.add_run(
                        f"{step.get('step', 'Cleaning Operation')}").bold = True
                    step_para.add_run('\n')

                    if step.get('description'):
                        desc_para = doc.add_paragraph()
                        desc_para.add_run('   Description: ').bold = True
                        desc_para.add_run(step['description'])

                    if step.get('details'):
                        details = step['details']
                        if isinstance(details, list):
                            for detail in details:
                                detail_para = doc.add_paragraph()
                                detail_para.add_run(f'   • {detail}')
                        else:
                            detail_para = doc.add_paragraph()
                            detail_para.add_run(f'   • {details}')
                else:
                    step_para.add_run(step).bold = True

                doc.add_paragraph()  # Add space between steps

        # Final Quality Assessment
        if cleaning_report:
            doc.add_heading('Post-Cleaning Quality Assessment', level=1)

            quality_para = doc.add_paragraph()
            quality_para.add_run('Quality Improvement Summary:\n').bold = True
            quality_para.add_run(
                f'• Initial Quality Score: {before_quality:.1f}%\n')
            quality_para.add_run(
                f'• Final Quality Score: {after_quality:.1f}%\n')
            quality_para.add_run(
                f'• Overall Improvement: {improvement:+.1f}%\n')
            quality_para.add_run(
                f'• Issues Successfully Resolved: {before_issues - after_issues}\n')
            quality_para.add_run(f'• Remaining Issues: {after_issues}\n\n')

            assessment_para = doc.add_paragraph()
            assessment_para.add_run('Assessment: ').bold = True
            if after_quality >= 95:
                assessment_para.add_run(
                    'Excellent - Data quality is now at optimal level for analysis and modeling.')
            elif after_quality >= 85:
                assessment_para.add_run(
                    'Good - Data quality is acceptable for most analytical purposes.')
            elif after_quality >= 70:
                assessment_para.add_run(
                    'Fair - Data quality has improved but may benefit from additional domain-specific cleaning.')
            else:
                assessment_para.add_run(
                    'Needs Attention - Consider additional data sources or manual review for critical applications.')

        # Recommendations
        doc.add_heading('Recommendations & Next Steps', level=1)

        rec_para = doc.add_paragraph()
        rec_para.add_run('Data Quality Recommendations:\n').bold = True
        rec_para.add_run(
            '• Implement regular data quality monitoring for ongoing data ingestion\n')
        rec_para.add_run(
            '• Consider implementing data validation rules at the source to prevent future quality issues\n')
        rec_para.add_run(
            '• Establish data governance policies for consistent data handling\n')
        rec_para.add_run(
            '• Regular automated quality checks should be scheduled\n')
        rec_para.add_run(
            '• Document data transformation rules for reproducibility\n\n')

        rec_para.add_run('Technical Recommendations:\n').bold = True
        rec_para.add_run('• Set up automated data validation pipelines\n')
        rec_para.add_run('• Implement exception handling for edge cases\n')
        rec_para.add_run('• Create data quality dashboards for monitoring\n')
        rec_para.add_run(
            '• Establish backup and recovery procedures for cleaned data\n')

        # Technical Appendix
        doc.add_heading('Technical Appendix', level=1)

        tech_para = doc.add_paragraph()
        tech_para.add_run('Quality Metrics Used:\n').bold = True
        tech_para.add_run('• Missing Value Percentage Analysis\n')
        tech_para.add_run('• Duplicate Record Detection and Removal\n')
        tech_para.add_run(
            '• Statistical Outlier Identification (IQR Method)\n')
        tech_para.add_run('• Data Type Consistency Validation\n')
        tech_para.add_run('• Text Format Standardization\n')
        tech_para.add_run('• Cross-field Validation Rules\n\n')

        tech_para.add_run('AI Algorithms Applied:\n').bold = True
        tech_para.add_run('• Pattern Recognition for Data Anomalies\n')
        tech_para.add_run('• Machine Learning-based Outlier Detection\n')
        tech_para.add_run('• Intelligent Missing Value Imputation\n')
        tech_para.add_run('• Automated Data Type Inference\n')
        tech_para.add_run('• Context-aware Text Cleaning\n')

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            'Generated by AI-Powered Data Cleaning Platform | Professional Data Quality Management')
        footer_run.italic = True

        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()  # Close the file handle before using it
        doc.save(temp_file.name)

        # Read the file content
        with open(temp_file.name, 'rb') as f:
            word_content = f.read()

        # Clean up
        os.unlink(temp_file.name)

        return word_content

    except ImportError:
        st.error("Word generation requires python-docx library. Please install it.")
        return None
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None


def create_upload_interface():
    """Create upload interface"""
    st.markdown("### Upload Your Dataset")

    # API Connection Status
    api_connected = check_api_connection()

    if api_connected:
        st.markdown("""
        <div class="status-success">
            <strong>System Online:</strong> Ready to process your data
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="status-error">
            <strong>System Offline:</strong> Backend service unavailable
        </div>
        """, unsafe_allow_html=True)
        return None

    # Upload Area
    st.markdown("""
    <div class="upload-container">
        <p>Supported formats: CSV, Excel (XLSX, XLS), pdf, png, JSON</p>
    </div>
    """, unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Choose file",
        type=['csv', 'xlsx', 'xls', 'json', 'pdf',
              'png', 'jpg', 'jpeg', 'html', 'txt'],
        label_visibility="collapsed"
    )

    return uploaded_file


def create_progress_tracker():
    """Create interactive progress tracker with clickable steps"""
    st.markdown("""
    <div class="progress-container">
        <h3 style="text-align: center; margin-bottom: 2rem; color: #2c3e50;">Processing Pipeline</h3>
        <div class="step-indicator">
    """, unsafe_allow_html=True)

    steps = ["Upload", "Analyze", "Clean", "Report"]
    step_descriptions = [
        "Upload your dataset",
        "Analyze data quality",
        "Clean and optimize",
        "View results"
    ]

    cols = st.columns(4)

    for i, (step, description) in enumerate(zip(steps, step_descriptions), 1):
        with cols[i-1]:
            # Determine step class
            if i < st.session_state.current_step:
                class_name = "step completed clickable"
                clickable = True
            elif i == st.session_state.current_step:
                class_name = "step active"
                clickable = False
            else:
                class_name = "step"
                clickable = False

            # Create clickable step if it's completed
            if clickable and i <= 2:  # Only allow clicking back to Upload or Analyze
                if st.button(f"{i}", key=f"step_{i}", help=f"Go back to: {description}"):
                    st.session_state.current_step = i
                    if i == 1:
                        # Reset session for new upload
                        for key in list(st.session_state.keys()):
                            if key.startswith(('file_', 'analysis_', 'cleaning_', 'original_', 'cleaned_')):
                                del st.session_state[key]
                        st.session_state.show_complete_report = False
                    st.rerun()
            else:
                # Non-clickable step indicator
                st.markdown(
                    f'<div class="{class_name}">{i}</div>', unsafe_allow_html=True)

            # Step label
            st.markdown(f"""
            <div style="text-align: center; margin-top: 0.5rem; font-size: 0.9rem; color: #666;">
                {step}
            </div>
            """, unsafe_allow_html=True)

    st.markdown('</div></div>', unsafe_allow_html=True)

    # Add step descriptions
    current_step_desc = step_descriptions[st.session_state.current_step - 1]
    st.markdown(f"""
    <div style="text-align: center; margin-top: 1rem; padding: 1rem; background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%); border-radius: 10px;">
        <strong>Current Step:</strong> {current_step_desc}
    </div>
    """, unsafe_allow_html=True)


def upload_file(file):
    """Upload file to backend"""
    global API_BASE_URL
    try:
        files = {"file": (file.name, file.getvalue(), file.type)}
        # Try current base, then fall back to 8000
        base_candidates = [API_BASE_URL,
                           "http://localhost:8000", "http://127.0.0.1:8000"]
        last_error = None
        for base in base_candidates:
            try:
                response = requests.post(
                    f"{base}/upload", files=files, timeout=600)
                if response.status_code == 200:
                    API_BASE_URL = base
                    st.session_state["api_base"] = base
                    return response.json()
                else:
                    last_error = f"{response.status_code}: {response.text}"
            except requests.Timeout:
                last_error = "Upload request timed out"
            except Exception as e:
                last_error = str(e)
        if last_error:
            st.error(f"Upload failed: {last_error}")
            return None
    except requests.Timeout:
        st.error(
            "Upload timed out. This file may be very large. Please try again; we increased the timeout to 10 minutes.")
        return None
    except Exception as e:
        st.error(f"Upload error: {str(e)}")
        return None


def _normalize_analysis_payload(payload: dict) -> dict:
    """Normalize different backend analysis payloads to a common schema used by the UI.

    Contract in -> out:
    - If payload already contains keys: 'shape', 'issues', 'quality_score', return as-is.
    - If payload is DataProfile-like (keys such as 'total_rows','total_columns','missing_values',
      'duplicates','outliers'), convert to:
        {
          'shape': {'rows': int, 'columns': int},
          'missing_by_column': {col: int},
          'issues': [ {type,column,count,percentage,description,severity} ],
          'quality_score': float,
        }
    """
    if not isinstance(payload, dict):
        return {}

    # If it already matches the autonomous_api schema, use it directly
    if 'shape' in payload and 'issues' in payload:
        return payload

    # DataProfile-style schema normalization
    if 'total_rows' in payload and 'total_columns' in payload:
        rows = int(payload.get('total_rows') or 0)
        cols = int(payload.get('total_columns') or 0)
        missing = payload.get('missing_values', {}) or {}
        dups = int(payload.get('duplicates') or 0)
        outliers = payload.get('outliers', {}) or {}

        issues = []
        quality_deductions = 0.0

        # Missing values issues
        for col, count in missing.items():
            try:
                mcount = int(count)
            except Exception:
                mcount = 0
            if mcount <= 0:
                continue
            pct = (mcount / rows * 100.0) if rows > 0 else 0.0
            severity = 'high' if pct > 30 else 'medium' if pct > 10 else 'low'
            issues.append({
                'type': 'missing_values',
                'column': col,
                'count': mcount,
                'percentage': round(pct, 2),
                'description': f"Column '{col}' has {mcount} missing values ({pct:.1f}%)",
                'severity': severity,
            })
            quality_deductions += min(pct * 0.5, 25)
            quality_deductions += min(pct * 0.5, 15)

        # Duplicate rows issue (aggregated)
        if dups > 0 and rows > 0:
            dup_pct = dups / rows * 100.0
            issues.append({
                'type': 'duplicates',
                'column': 'all_columns',
                'count': dups,
                'percentage': round(dup_pct, 2),
                'description': f"Found {dups} duplicate rows ({dup_pct:.1f}% of data)",
                'severity': 'high' if dup_pct > 10 else 'medium' if dup_pct > 5 else 'low',
            })
            quality_deductions += min(dup_pct * 0.8, 20)

        # Outliers per numeric column
        for col, count in outliers.items():
            try:
                ocount = int(count)
            except Exception:
                ocount = 0
            if ocount <= 0 or rows <= 0:
                continue
            pct = (ocount / rows * 100.0)
            issues.append({
                'type': 'outliers',
                'column': col,
                'count': ocount,
                'percentage': round(pct, 2),
                'description': f"Column '{col}' has {ocount} statistical outliers ({pct:.1f}%)",
                'severity': 'medium' if pct > 5 else 'low',
            })
            quality_deductions += min(pct * 0.3, 10)

        quality_score = max(0.0, 100.0 - quality_deductions)

        normalized = {
            'shape': {'rows': rows, 'columns': cols},
            'missing_by_column': {k: int(v) for k, v in (missing or {}).items()},
            'issues': issues,
            'quality_score': round(quality_score, 1),
        }
        # Keep raw details for possible future charts
        if 'data_types' in payload:
            normalized['data_types'] = payload.get('data_types')
        if 'column_info' in payload:
            normalized['column_info'] = payload.get('column_info')
        return normalized

    # Unknown structure
    return payload


def get_analysis(file_id):
    """Get analysis from backend with smart route and port fallback.

    Tries both base ports (current base, 8000 fallback) and these routes per base:
    - /analyze/{file_id}
    - /profile/{file_id}
    - /api/v1/profile/{file_id}
    Updates the global API_BASE_URL when a working base is found, and normalizes
    the payload so the UI renders consistently across backends.
    """
    global API_BASE_URL
    base_candidates = []
    # Prefer the current base first
    base_candidates.append(API_BASE_URL)
    # Add 8000 fallback explicitly
    if "8000" not in API_BASE_URL:
        base_candidates.append("http://localhost:8000")
        base_candidates.append("http://127.0.0.1:8000")
    else:
        base_candidates.append("http://localhost:8003")
        base_candidates.append("http://127.0.0.1:8003")

    last_error = None
    for base in base_candidates:
        routes = [
            f"{base}/analyze/{file_id}",
            f"{base}/profile/{file_id}",
            f"{base}/api/v1/profile/{file_id}",
        ]
        for url in routes:
            try:
                response = requests.get(url, timeout=60)
                if response.status_code == 200:
                    API_BASE_URL = base  # Remember the working base
                    st.session_state["api_base"] = base
                    try:
                        raw = response.json()
                    except Exception:
                        raw = None
                    return _normalize_analysis_payload(raw or {})
                else:
                    last_error = f"{response.status_code}: {response.text}"
            except requests.Timeout:
                last_error = "Request timed out"
            except Exception as e:
                last_error = str(e)
    if last_error:
        st.error(f"Analysis failed: {last_error}")
    return None


def clean_data(file_id):
    """Clean data using backend with base fallback"""
    global API_BASE_URL
    base_candidates = [API_BASE_URL,
                       "http://localhost:8000", "http://127.0.0.1:8000"]
    last_error = None
    for base in base_candidates:
        try:
            response = requests.post(f"{base}/clean/{file_id}", timeout=600)
            if response.status_code == 200:
                API_BASE_URL = base
                st.session_state["api_base"] = base
                cleaning_result = response.json()

                # Try to fetch previews (best effort; may not exist on all backends)
                try:
                    original_response = requests.get(
                        f"{base}/data/{file_id}/original", timeout=30)
                    if original_response.status_code == 200:
                        original_data_json = original_response.json()
                        st.session_state.original_data = pd.DataFrame(
                            original_data_json.get('data', []))

                    cleaned_response = requests.get(
                        f"{base}/data/{file_id}/cleaned", timeout=30)
                    if cleaned_response.status_code == 200:
                        cleaned_data_json = cleaned_response.json()
                        st.session_state.cleaned_data = pd.DataFrame(
                            cleaned_data_json.get('data', []))
                except Exception as e:
                    st.warning(f"Could not fetch data for display: {e}")

                return cleaning_result
            else:
                last_error = f"{response.status_code}: {response.text}"
        except requests.Timeout:
            last_error = "Cleaning request timed out"
        except Exception as e:
            last_error = str(e)
    if last_error:
        st.error(f"Cleaning failed: {last_error}")
    return None


def create_comprehensive_analysis_dashboard(analysis_data):
    """Create detailed analysis dashboard with useful visualizations"""
    st.markdown('<div class="section-header">Data Quality Analysis Report</div>',
                unsafe_allow_html=True)

    # Key Metrics
    quality_score = analysis_data.get('quality_score', 0)
    total_issues = len(analysis_data.get('issues', []))
    rows = analysis_data.get('shape', {}).get('rows', 0)
    cols = analysis_data.get('shape', {}).get('columns', 0)

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.markdown(f"""
        <div class="metric-container">
            <div class="metric-title">Overall Quality</div>
            <div class="metric-value">{quality_score:.1f}%</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="metric-container">
            <div class="metric-title">Issues Detected</div>
            <div class="metric-value">{total_issues}</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown(f"""
        <div class="metric-container">
            <div class="metric-title">Records</div>
            <div class="metric-value">{rows:,}</div>
        </div>
        """, unsafe_allow_html=True)

    with col4:
        st.markdown(f"""
        <div class="metric-container">
            <div class="metric-title">Columns</div>
            <div class="metric-value">{cols}</div>
        </div>
        """, unsafe_allow_html=True)

    # Useful Visualizations with enhanced graphics
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<div class="chart-container scale-in">',
                    unsafe_allow_html=True)
        st.markdown(
            '<div class="chart-title">Issues by Severity Level</div>', unsafe_allow_html=True)

        if analysis_data.get('issues'):
            issues_df = pd.DataFrame(analysis_data['issues'])
            severity_counts = issues_df['severity'].value_counts()

            # Enhanced color scheme with gradients
            colors = ['#e74c3c', '#f39c12', '#f1c40f']

            fig = go.Figure(data=[go.Pie(
                labels=severity_counts.index,
                values=severity_counts.values,
                hole=.4,
                marker=dict(
                    colors=colors,
                    line=dict(color='#FFFFFF', width=3)
                ),
                textinfo='label+percent',
                textposition='auto',
                textfont=dict(size=18, color='white', family='Roboto'),
                hovertemplate='<b>%{label}</b><br>Count: %{value}<br>Percentage: %{percent}<extra></extra>',
                pull=[0.1 if sev == 'high' else 0 for sev in severity_counts.index]
            )])

            fig.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font=dict(size=20, color='#2c3e50', family='Roboto'),
                title=dict(
                    text="Issues by Severity",
                    font=dict(size=24, color='#2c3e50'),
                    x=0.5
                ),
                showlegend=True,
                legend=dict(
                    orientation="v",
                    yanchor="middle",
                    y=0.5,
                    xanchor="left",
                    x=1.05
                ),
                margin=dict(t=50, b=50, l=50, r=100),
                annotations=[dict(text='Issues', x=0.5, y=0.5,
                                  font_size=20, showarrow=False, font_color='#2c3e50')]
            )

            st.plotly_chart(fig, use_container_width=True)
        else:
            st.markdown("""
            <div style="text-align: center; padding: 3rem; color: #27ae60;">
                <div style="font-size: 4rem; margin-bottom: 1rem;">PERFECT</div>
                <h3>No Issues Detected!</h3>
                <p>Your data quality is excellent</p>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="chart-container scale-in">',
                    unsafe_allow_html=True)
        st.markdown(
            '<div class="chart-title">Missing Values by Column</div>', unsafe_allow_html=True)

        if analysis_data.get('missing_by_column'):
            missing_data = analysis_data['missing_by_column']
            if missing_data:
                missing_df = pd.DataFrame(list(missing_data.items()),
                                          columns=['Column', 'Missing_Count'])
                missing_df['Missing_Percentage'] = (
                    missing_df['Missing_Count'] / rows) * 100
                missing_df = missing_df.sort_values(
                    'Missing_Percentage', ascending=True)

                # Enhanced bar chart with gradient colors
                fig = go.Figure(data=[go.Bar(
                    y=missing_df['Column'],
                    x=missing_df['Missing_Percentage'],
                    orientation='h',
                    marker=dict(
                        color=missing_df['Missing_Percentage'],
                        colorscale='Reds',
                        colorbar=dict(
                            title=dict(text="Missing %", font=dict(size=16)),
                            tickmode="linear",
                            tick0=0,
                            dtick=10,
                            tickfont=dict(size=14)
                        ),
                        line=dict(color='rgba(255,255,255,0.6)', width=1)
                    ),
                    hovertemplate='<b>%{y}</b><br>Missing: %{x:.1f}%<br>Count: ' +
                    missing_df['Missing_Count'].astype(
                        str) + '<extra></extra>',
                    text=missing_df['Missing_Percentage'].round(
                        1).astype(str) + '%',
                    textposition='auto'
                )])

                fig.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    font=dict(size=18, color='#2c3e50'),
                    title=dict(
                        text="Missing Data Analysis",
                        font=dict(size=24, color='#2c3e50'),
                        x=0.5
                    ),
                    xaxis=dict(
                        title=dict(text="Missing Data Percentage (%)",
                                   font=dict(size=18)),
                        showgrid=True,
                        gridcolor='rgba(0,0,0,0.1)',
                        range=[0, max(missing_df['Missing_Percentage']) * 1.1],
                        tickfont=dict(size=16)
                    ),
                    yaxis=dict(
                        title=dict(text="Columns", font=dict(size=18)),
                        showgrid=False,
                        tickfont=dict(size=16)
                    ),
                    margin=dict(t=70, b=60, l=120, r=60),
                    height=500
                )

                st.plotly_chart(fig, use_container_width=True)
            else:
                st.markdown("""
                <div style="text-align: center; padding: 3rem; color: #27ae60;">
                    <div style="font-size: 4rem; margin-bottom: 1rem;">COMPLETE</div>
                    <h3>Complete Data!</h3>
                    <p>No missing values detected</p>
                </div>
                """, unsafe_allow_html=True)

        # Detailed Issues Report (Web-based, not downloadable)
        st.markdown('</div>', unsafe_allow_html=True)
    if analysis_data.get('issues'):
        st.markdown('<div class="report-section">', unsafe_allow_html=True)
        st.markdown(
            '<div class="section-header">Detailed Issues Report</div>', unsafe_allow_html=True)

        # Make issues collapsible with expander
        with st.expander("Click to view detailed issues breakdown", expanded=False):
            issues_df = pd.DataFrame(analysis_data['issues'])

            # Group issues by severity for better display
            for severity in ['high', 'medium', 'low']:
                severity_issues = issues_df[issues_df['severity'] == severity]
                if len(severity_issues) > 0:
                    severity_label = {'high': 'HIGH',
                                      'medium': 'MEDIUM', 'low': 'LOW'}
                    st.markdown(
                        f"### {severity_label[severity]} Severity Issues ({len(severity_issues)})")

                    for _, issue in severity_issues.iterrows():
                        issue_class = f"issue-item issue-{severity}"
                        st.markdown(f"""
                        <div class="{issue_class}">
                            <strong>Column:</strong> {issue['column']}<br>
                            <strong>Issue:</strong> {issue['type'].replace('_', ' ').title()}<br>
                            <strong>Count:</strong> {issue['count']} ({issue['percentage']}%)<br>
                            <strong>Description:</strong> {issue['description']}
                        </div>
                        """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)


def create_comprehensive_cleaning_dashboard(cleaning_report):
    """Create detailed cleaning results dashboard with web-based reports"""
    st.markdown('<div class="section-header">Data Cleaning Results</div>',
                unsafe_allow_html=True)

    # Before/After Comparison
    before_quality = cleaning_report.get(
        'before_quality', {}).get('quality_score', 0)
    after_quality = cleaning_report.get(
        'after_quality', {}).get('quality_score', 0)
    improvement = after_quality - before_quality

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(f"""
        <div class="metric-container" style="border-left-color: #e74c3c;">
            <div class="metric-title">Before Cleaning</div>
            <div class="metric-value">{before_quality:.1f}%</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="metric-container" style="border-left-color: #27ae60;">
            <div class="metric-title">After Cleaning</div>
            <div class="metric-value">{after_quality:.1f}%</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        improvement_color = "#27ae60" if improvement >= 0 else "#e74c3c"
        improvement_symbol = "+" if improvement >= 0 else ""
        st.markdown(f"""
        <div class="metric-container" style="border-left-color: {improvement_color};">
            <div class="metric-title">Improvement</div>
            <div class="metric-value" style="color: {improvement_color};">{improvement_symbol}{improvement:.1f}%</div>
        </div>
        """, unsafe_allow_html=True)

    # Quality Improvement Visualization with enhanced graphics
    st.markdown('<div class="chart-container scale-in">',
                unsafe_allow_html=True)
    st.markdown('<div class="chart-title">Quality Score Improvement</div>',
                unsafe_allow_html=True)

    # Create a more sophisticated chart
    fig = go.Figure()

    # Add background gradient effect
    fig.add_shape(
        type="rect",
        x0=-0.5, y0=0, x1=1.5, y1=100,
        fillcolor="rgba(52, 152, 219, 0.1)",
        layer="below",
        line_width=0,
    )

    # Before bar with gradient
    fig.add_trace(go.Bar(
        x=['Before Cleaning'],
        y=[before_quality],
        name='Before Cleaning',
        marker=dict(
            color=before_quality,
            colorscale=[[0, '#e74c3c'], [0.5, '#f39c12'], [1, '#f1c40f']],
            cmin=0,
            cmax=100,
            line=dict(color='rgba(255,255,255,0.6)', width=2)
        ),
        text=[f'{before_quality:.1f}%'],
        textposition='auto',
        textfont=dict(color='white', size=16, family='Roboto'),
        width=0.4,
        hovertemplate='<b>Before Cleaning</b><br>Quality Score: %{y:.1f}%<extra></extra>'
    ))

    # After bar with gradient
    fig.add_trace(go.Bar(
        x=['After Cleaning'],
        y=[after_quality],
        name='After Cleaning',
        marker=dict(
            color=after_quality,
            colorscale=[[0, '#e74c3c'], [0.7, '#f39c12'],
                        [0.85, '#27ae60'], [1, '#2ecc71']],
            cmin=0,
            cmax=100,
            line=dict(color='rgba(255,255,255,0.6)', width=2)
        ),
        text=[f'{after_quality:.1f}%'],
        textposition='auto',
        textfont=dict(color='white', size=16, family='Roboto'),
        width=0.4,
        hovertemplate='<b>After Cleaning</b><br>Quality Score: %{y:.1f}%<extra></extra>'
    ))

    # Add improvement arrow if there's improvement
    if improvement > 0:
        fig.add_annotation(
            x=0.5,
            y=max(before_quality, after_quality) + 10,
            text=f"+{improvement:.1f}%",
            showarrow=True,
            arrowhead=2,
            arrowsize=1,
            arrowwidth=3,
            arrowcolor="#27ae60",
            font=dict(size=18, color="#27ae60", family='Roboto'),
            bgcolor="rgba(255,255,255,0.8)",
            bordercolor="#27ae60",
            borderwidth=2
        )

    fig.update_layout(
        title=dict(
            text="Data Quality Comparison",
            font=dict(size=26, color='#2c3e50'),
            x=0.5
        ),
        yaxis=dict(
            title=dict(text="Quality Score (%)", font=dict(size=20)),
            range=[0, 105],
            showgrid=True,
            gridcolor='rgba(0,0,0,0.1)',
            tickfont=dict(size=18, color='#2c3e50')
        ),
        xaxis=dict(
            title=dict(text="Data State", font=dict(size=20)),
            tickfont=dict(size=18, color='#2c3e50')
        ),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(size=18, color='#2c3e50', family='Roboto'),
        showlegend=False,
        margin=dict(t=80, b=60, l=80, r=60),
        bargap=0.3
    )

    st.plotly_chart(fig, use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Data Comparison (Show actual data on web)
    st.markdown('<div class="report-section">', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Data Comparison</div>',
                unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Original Data Sample:**")
        if st.session_state.original_data is not None and not st.session_state.original_data.empty:
            st.markdown('<div class="data-preview">', unsafe_allow_html=True)
            st.dataframe(st.session_state.original_data.head(
                10), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("Original data not available for preview")

    with col2:
        st.markdown("**Cleaned Data Sample:**")
        if st.session_state.cleaned_data is not None and not st.session_state.cleaned_data.empty:
            st.markdown('<div class="data-preview">', unsafe_allow_html=True)
            st.dataframe(st.session_state.cleaned_data.head(
                10), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("Cleaned data not available for preview")

    st.markdown('</div>', unsafe_allow_html=True)

    # Export Options
    st.markdown('<div class="section-header">Export Options</div>',
                unsafe_allow_html=True)

    # Main action buttons layout
    col1, col2, col3 = st.columns([2, 2, 2])

    with col1:
        # Download cleaned dataset only
        st.markdown("**Download Cleaned Data:**")
        try:
            response = requests.get(
                f"{API_BASE_URL}/download/{st.session_state.file_id}/cleaned")
            if response.status_code == 200:
                st.download_button(
                    label="Download CSV",
                    data=response.content,
                    file_name=f"cleaned_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    key="csv_download",
                    type="primary",
                    use_container_width=True
                )
            else:
                st.error("Unable to download cleaned data")
        except Exception as e:
            st.error(f"Download error: {e}")

    with col2:
        # View report section
        st.markdown("**View Report:**")
        st.markdown('<br>', unsafe_allow_html=True)
        if st.button("View Complete Report", key="view_report", use_container_width=True):
            st.session_state.show_complete_report = True
            st.rerun()

    with col3:
        # New dataset section
        st.markdown("**New Analysis:**")
        st.markdown('<br>', unsafe_allow_html=True)
        if st.button("Process New Dataset", key="new_file", use_container_width=True):
            # Reset session state
            for key in list(st.session_state.keys()):
                if key.startswith(('file_', 'analysis_', 'cleaning_', 'original_', 'cleaned_')):
                    del st.session_state[key]
            st.session_state.current_step = 1
            st.rerun()


def create_complete_detailed_report(cleaning_report):
    """Create comprehensive detailed report with all information"""
    st.markdown('<div class="section-header">Complete Data Cleaning Report</div>',
                unsafe_allow_html=True)

    # Report Header
    st.markdown('<div class="report-section">', unsafe_allow_html=True)
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.markdown(f"""
    <div style="background: #f8f9fa; padding: 1rem; border-radius: 8px; margin-bottom: 1rem;">
        <h3 style="margin: 0; color: #2c3e50;">Data Quality Assessment & Cleaning Report</h3>
        <p style="margin: 0.5rem 0 0 0; color: #666;">Generated on: {current_time}</p>
    </div>
    """, unsafe_allow_html=True)

    # Executive Summary
    st.markdown("### Executive Summary")
    before_quality = cleaning_report.get(
        'before_quality', {}).get('quality_score', 0)
    after_quality = cleaning_report.get(
        'after_quality', {}).get('quality_score', 0)
    improvement = after_quality - before_quality

    before_issues = len(cleaning_report.get(
        'before_quality', {}).get('issues', []))
    after_issues = len(cleaning_report.get(
        'after_quality', {}).get('issues', []))

    before_shape = cleaning_report.get('shape_change', {}).get('before', {})
    after_shape = cleaning_report.get('shape_change', {}).get('after', {})

    st.markdown(f"""
    - **Overall Quality Improvement**: {before_quality:.1f}% → {after_quality:.1f}% ({improvement:+.1f}%)
    - **Issues Resolved**: {before_issues} → {after_issues} ({before_issues - after_issues} issues fixed)
    - **Data Shape**: {before_shape.get('rows', 0):,} rows × {before_shape.get('columns', 0)} columns → {after_shape.get('rows', 0):,} rows × {after_shape.get('columns', 0)} columns
    - **Data Integrity**: Maintained with improved consistency and completeness
    """)

    # Original Issues Found
    if cleaning_report.get('before_quality', {}).get('issues'):
        st.markdown("### Original Data Quality Issues")
        issues_df = pd.DataFrame(cleaning_report['before_quality']['issues'])

        for severity in ['high', 'medium', 'low']:
            severity_issues = issues_df[issues_df['severity'] == severity]
            if len(severity_issues) > 0:
                st.markdown(
                    f"**{severity.upper()} Priority Issues ({len(severity_issues)}):**")
                for _, issue in severity_issues.iterrows():
                    st.markdown(
                        f"- **{issue['column']}**: {issue['type'].replace('_', ' ').title()} ({issue['count']} instances, {issue['percentage']}%)")
                    st.markdown(f"  *{issue['description']}*")

    # Applied Cleaning Operations
    if cleaning_report.get('steps_applied'):
        st.markdown("### Applied Cleaning Operations")
        steps = cleaning_report['steps_applied']

        for i, step in enumerate(steps, 1):
            if isinstance(step, dict):
                st.markdown(
                    f"**{i}. {step.get('step', 'Cleaning Operation')}**")
                if step.get('description'):
                    st.markdown(f"   - Description: {step['description']}")
                if step.get('details'):
                    details = step['details']
                    if isinstance(details, list):
                        for detail in details:
                            st.markdown(f"   - {detail}")
                    else:
                        st.markdown(f"   - {details}")
            else:
                st.markdown(f"**{i}. {step}**")

    # Final Data Quality Assessment
    if cleaning_report.get('after_quality', {}).get('issues'):
        st.markdown("### Remaining Issues (Post-Cleaning)")
        remaining_issues = cleaning_report['after_quality']['issues']
        if remaining_issues:
            issues_df = pd.DataFrame(remaining_issues)
            for _, issue in issues_df.iterrows():
                st.markdown(
                    f"- **{issue['column']}**: {issue['type'].replace('_', ' ').title()} ({issue['count']} instances)")
        else:
            st.success("No remaining data quality issues detected!")

    # Recommendations
    st.markdown("### Recommendations")
    if after_quality >= 95:
        st.markdown(
            "- **Excellent**: Data quality is now at an optimal level for analysis and modeling")
    elif after_quality >= 85:
        st.markdown(
            "- **Good**: Data quality is acceptable, minor issues may remain but won't significantly impact analysis")
    elif after_quality >= 70:
        st.markdown(
            "- **Fair**: Data quality has improved but may benefit from additional domain-specific cleaning")
    else:
        st.markdown(
            "- **Needs Attention**: Consider additional data sources or manual review for critical applications")

    st.markdown(
        "- Regular data quality monitoring is recommended for ongoing data ingestion")
    st.markdown(
        "- Consider implementing data validation rules at the source to prevent future quality issues")

    # Technical Details
    with st.expander("Technical Details"):
        st.markdown("**Data Types:**")
        if cleaning_report.get('before_quality', {}).get('shape'):
            original_shape = cleaning_report['before_quality']['shape']
            st.markdown(
                f"- Original dataset: {original_shape['rows']:,} rows × {original_shape['columns']} columns")

        if cleaning_report.get('after_quality', {}).get('shape'):
            final_shape = cleaning_report['after_quality']['shape']
            st.markdown(
                f"- Final dataset: {final_shape['rows']:,} rows × {final_shape['columns']} columns")

        st.markdown("**Quality Metrics Used:**")
        st.markdown("- Missing value percentage")
        st.markdown("- Duplicate record detection")
        st.markdown("- Outlier identification using IQR method")
        st.markdown("- Data type consistency validation")
        st.markdown("- Text formatting standardization")

    # Download Reports Section
    st.markdown("### Download Reports")

    # Center the download section
    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        # Format selection with availability check
        available_formats = []
        if REPORTLAB_AVAILABLE:
            available_formats.append("PDF Report")
        if DOCX_AVAILABLE:
            available_formats.append("Word Document")

        if not available_formats:
            st.warning(
                "⚠️ Report generation packages not available. Please install 'reportlab' for PDF reports and/or 'python-docx' for Word documents.")
            st.info(
                "Run: `pip install reportlab python-docx` to enable report downloads")
        else:
            report_format = st.selectbox(
                "Select report format:",
                available_formats,
                key="complete_report_format"
            )

            # Single-click download button
            st.markdown("<br>", unsafe_allow_html=True)

            try:
                if report_format == "PDF Report" and REPORTLAB_AVAILABLE:
                    # Get analysis data for PDF generation
                    analysis_data = st.session_state.get('analysis_data', {})
                    pdf_content = generate_pdf_report(
                        analysis_data, cleaning_report)
                    if pdf_content:
                        st.download_button(
                            label="Download PDF Report",
                            data=pdf_content,
                            file_name=f"complete_data_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            key="complete_pdf_download",
                            type="primary",
                            use_container_width=True
                        )

                elif report_format == "Word Document" and DOCX_AVAILABLE:
                    # Get analysis data for Word generation
                    analysis_data = st.session_state.get('analysis_data', {})
                    word_content = generate_word_report(
                        analysis_data, cleaning_report)
                    if word_content:
                        st.download_button(
                            label="Download Word Report",
                            data=word_content,
                            file_name=f"complete_data_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="complete_word_download",
                            type="primary",
                            use_container_width=True
                        )
            except Exception as e:
                st.error(f"Error generating report: {e}")

    st.markdown('</div>', unsafe_allow_html=True)


def main():
    """Main application with enhanced animations and interactions"""
    create_header()

    # Add smooth transition effects
    st.markdown('<div class="fade-in">', unsafe_allow_html=True)

    # Main content based on current step
    if st.session_state.current_step == 1:
        # Upload page with animations
        st.markdown('<div class="slide-in-left">', unsafe_allow_html=True)
        uploaded_file = create_upload_interface()
        st.markdown('</div>', unsafe_allow_html=True)

        if uploaded_file:
            # Enhanced loading animation
            st.markdown("""
            <div style="text-align: center; padding: 2rem;">
                <div class="loading-spinner"></div>
                <h3 style="color: #3498db; margin-top: 1rem;">Processing your data...</h3>
                <p style="color: #666;">Analyzing data quality and structure</p>
            </div>
            """, unsafe_allow_html=True)

            with st.spinner("Processing and analyzing your data..."):
                upload_result = upload_file(uploaded_file)

                if upload_result:
                    st.session_state.file_id = upload_result['file_id']
                    # Prefer analysis from upload response; fallback to on-demand analysis
                    initial_analysis = upload_result.get('analysis')
                    if initial_analysis:
                        st.session_state.analysis_data = _normalize_analysis_payload(
                            initial_analysis)
                    else:
                        st.session_state.analysis_data = None
                    if not st.session_state.analysis_data and st.session_state.file_id:
                        # Some backends don't return analysis on upload; fetch it now
                        st.session_state.analysis_data = get_analysis(
                            st.session_state.file_id)
                    st.session_state.current_step = 2

                    # Success animation
                    st.markdown("""
                    <div style="text-align: center; padding: 2rem; animation: scaleIn 0.8s ease-out;">
                        <div style="font-size: 4rem; color: #27ae60; margin-bottom: 1rem;">SUCCESS</div>
                        <h3 style="color: #27ae60;">Upload Successful!</h3>
                        <p>Data analysis completed successfully</p>
                    </div>
                    """, unsafe_allow_html=True)

                    time.sleep(2)
                    st.rerun()

    elif st.session_state.current_step >= 2:
        # Processing interface with step tracker
        create_progress_tracker()

        if st.session_state.current_step == 2:
            # Analysis step with enhanced animations
            st.markdown('<div class="slide-in-right">', unsafe_allow_html=True)
            # If analysis isn't available yet, try to fetch it now (supports both backends)
            if not st.session_state.analysis_data and st.session_state.file_id:
                with st.spinner("Generating analysis report..."):
                    st.session_state.analysis_data = get_analysis(
                        st.session_state.file_id)
            if st.session_state.analysis_data:
                create_comprehensive_analysis_dashboard(
                    st.session_state.analysis_data)

                # Enhanced action button with animation
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.markdown(
                        '<div style="text-align: center; margin: 2rem 0;">', unsafe_allow_html=True)
                    if st.button("Start Data Cleaning", key="start_cleaning", type="primary"):
                        st.session_state.current_step = 3
                        st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.error(
                    "Unable to generate analysis. Please verify the backend is running and try again.")
            st.markdown('</div>', unsafe_allow_html=True)

        elif st.session_state.current_step == 3:
            # Cleaning step with enhanced progress animation
            st.markdown('<div class="slide-in-left">', unsafe_allow_html=True)
            st.markdown("### AI-Powered Data Cleaning in Progress")

            if st.session_state.cleaning_report is None:
                # Enhanced progress animation
                progress_placeholder = st.empty()

                with progress_placeholder.container():
                    st.markdown("""
                    <div style="text-align: center; padding: 3rem;">
                        <div class="loading-spinner" style="margin: 0 auto 2rem auto;"></div>
                        <h3 style="color: #3498db;">AI Cleaning in Progress...</h3>
                        <div style="background: #f8f9fa; border-radius: 10px; padding: 1rem; margin: 2rem 0;">
                            <div style="display: flex; justify-content: space-between; margin-bottom: 0.5rem;">
                                <span>Analyzing patterns...</span>
                                <span style="color: #27ae60;">DONE</span>
                            </div>
                            <div style="display: flex; justify-content: space-between; margin-bottom: 0.5rem;">
                                <span>Applying cleaning algorithms...</span>
                                <span style="color: #3498db;">⟳</span>
                            </div>
                            <div style="display: flex; justify-content: space-between;">
                                <span>Optimizing data quality...</span>
                                <span style="color: #95a5a6;">○</span>
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                with st.spinner("Applying intelligent cleaning algorithms..."):
                    cleaning_result = clean_data(st.session_state.file_id)

                    if cleaning_result:
                        st.session_state.cleaning_report = cleaning_result
                        st.session_state.current_step = 4

                        progress_placeholder.empty()

                        # Success animation
                        st.markdown("""
                        <div style="text-align: center; padding: 2rem; animation: scaleIn 0.8s ease-out;">
                            <div style="font-size: 4rem; color: #27ae60; margin-bottom: 1rem;">COMPLETE</div>
                            <h3 style="color: #27ae60;">Cleaning Complete!</h3>
                            <p>Your data has been optimized and cleaned</p>
                        </div>
                        """, unsafe_allow_html=True)

                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error("Cleaning failed. Please try again.")
            st.markdown('</div>', unsafe_allow_html=True)

        elif st.session_state.current_step == 4:
            # Results step with enhanced presentation
            st.markdown('<div class="scale-in">', unsafe_allow_html=True)
            if st.session_state.show_complete_report:
                # Show complete detailed report
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button("← Back to Summary", key="back_to_summary"):
                        st.session_state.show_complete_report = False
                        st.rerun()

                if st.session_state.cleaning_report:
                    create_complete_detailed_report(
                        st.session_state.cleaning_report)
            else:
                # Show summary dashboard
                if st.session_state.cleaning_report:
                    create_comprehensive_cleaning_dashboard(
                        st.session_state.cleaning_report)
            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()
